### dilekce_service.py (Yeni Modül)

from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

# Türkçe büyük harf dönüşümü (basvuru.py’den birebir)
tr_cevir = str.maketrans("çğiöşü", "ÇĞİÖŞÜ")
def buyuk_harf_tr(text):
    return text.translate(tr_cevir).upper()

# Orijinal şablon metni (kesinlikle değişmeden kopyalanmıştır)
deger_kaybi_sablon = """{sigorta_sirketi} A.Ş.’YE

TALEPTE BULUNAN: {ad_soyad} TC: {tc_no}
VEKİLİ                         : Av. Fatih DİŞÇİ
\t\t\tİsmet Kaptan Mah. Yeni Asır İş Merkezi No:3 Kat:6 Ofis No:606
\t\t\tKonak/İZMİR

AÇIKLAMALAR        :
{kaza_tarihi} tarihinde şirketinize Zorunlu Mali Sorumluluk Sigortası (Z.M.M.S.) ile sigortalı {karsi_plaka} plaka sayılı araç, %100 KUSURLA, müvekkile ait {muvekkil_plaka} plaka sayılı araç ile çift taraflı trafik kazasına karışmak suretiyle müvekkilin aracının hasarlanmasına sebep olmuştur.

Ekte yer alan trafik kazası tespit tutanağına göre kazada {karsi_plaka} plaka sayılı araç %100 KUSURLU, müvekkile ait {muvekkil_plaka} plaka sayılı araç ise KUSURSUZ durumdadır.

Müvekkile ait araç {arac_modeli} model ve orjinal bir araçken, meydana gelen çift taraflı trafik kazası sebebiyle çeşitli yerlerinden darbeler alarak kapsamlı bir tamir, onarım, bakım sürecinden geçmiştir ve bu sebeple aracın ikinci el piyasa rayiç değerinde ciddi bir düşüş yaşanması kaçınılmazdır. Karayolları Trafik Kanunu Zorunlu Mali Sorumluluk Sigortası Genel Şartları’nın A.5. maddesi kapsamında değer kaybı maddi zararlar teminatının içinde sayılmıştır.

{kaza_tarihi} tarihinde yaşanan çift taraflı trafik kazası neticesinde müvekkil aracının ikinci el rayiç değeri {deger_kaybi} TL düşmüştür ve bu değer kaybının şirketinizce karşılanması gerektiği hususu kaçınılmazdır.

Ekte sunulan kaza tespit tutanağından da anlaşıldığı üzere mevcut kazaya sigortalınız tarafından sebebiyet verilmiş ve müvekkil maddi kayba uğratılmıştır. Bu sebeple; müvekkilin aracında meydana gelen hasardan ötürü aracın birçok önemli parçası işlem ve onarım görmüştür. Bu nedenle hasar bedelinin, orijinal parça farkının, iskontosuz hasar bedelinin ve ödenmesi gereken KDV ile diğer alacak kalemlerinin tazmin edilmesi gerektiği açıktır.

Yapılan emsal fiyat araştırmaları sonrası, piyasa koşullarına göre belirlenen {bakiye_hasar} TL bakiye hasar tazminatının, tarafınızca müvekkilime ödenmesi gerekmektedir.

SONUÇ VE İSTEM      :
Yukarıda arz ve izah edildiği üzere; müvekkilin değer kaybı tutarının {deger_kaybi} TL, bakiye hasar tutarının {bakiye_hasar} TL olduğundan, prim aşımı da göz önüne alındığında toplam {toplam} TL tutarındaki zararın ödenmesine ve yargılama giderleri ile vekalet ücretinin davalıya yükletilmesine karar verilmesini saygıyla arz ve talep ederiz.

BAŞVURAN VEKİLİ
Av. Fatih DİŞÇİ

EK: 
1- Kaza Tespit Tutanağı 2- Araç Ruhsat Fotokopisi 3- Eksper Raporu 4- Kaza Fotoğrafları 5- Vekâletname

VEKİL HESAP BİLGİLERİ: Türk Ekonomi Bankası TR45 0003 2000 0000 0106 9135 78  Hesap Sahibi: Fatih DİŞÇİ

İLETİŞİM BİLGİLERİ: fatihdisci@outlook.com  (0 507) 724 77 35"""


def create_insurance_docx(data: dict, output_path: str):
    # Document ve font ayarı
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Şablonu doldur
    doldurulmus = deger_kaybi_sablon.format(
        sigorta_sirketi = buyuk_harf_tr(data['Sigorta Şirketi']),
        ad_soyad        = buyuk_harf_tr(data['Ad Soyad']),
        tc_no           = data['TC No'],
        kaza_tarihi     = data['Kaza Tarihi'],
        muvekkil_plaka  = buyuk_harf_tr(data['Müvekkil Plaka']),
        karsi_plaka     = buyuk_harf_tr(data['Karşı Plaka']),
        arac_modeli     = data['Araç Modeli'],
        deger_kaybi     = data['Değer Kaybı (₺)'],
        bakiye_hasar    = data['Bakiye Hasar (₺)'],
        toplam          = data['Toplam (₺)']
    )

    for paragraf in doldurulmus.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        text = paragraf.strip()
        if text.upper().startswith("TALEPTE BULUNAN:"):
            run1 = p.add_run("TALEPTE BULUNAN:")
            run1.bold = True
            run1.underline = True
            run2 = p.add_run(" " + text[len("TALEPTE BULUNAN:"):].strip())
        elif text.upper().startswith("VEKİLİ                         :"):
            run1 = p.add_run("VEKİLİ                         :")
            run1.bold = True
            run1.underline = True
            run2 = p.add_run(" " + text[len("VEKİLİ                         :"):].strip())
        elif text.upper().endswith("A.Ş.’YE"):
            run = p.add_run(text.upper())
            run.bold = True
            run.underline = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif text.upper() in ["AÇIKLAMALAR        :", "SONUÇ VE İSTEM      :", "EK:", "VEKİL HESAP BİLGİLERİ:", "İLETİŞİM BİLGİLERİ:"]:
            run = p.add_run(text.upper())
            run.bold = True
            run.underline = True
        elif text.upper() == "BAŞVURAN VEKİLİ":
            run = p.add_run(text.upper())
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif text.startswith("Av. Fatih DİŞÇİ"):
            run = p.add_run(text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif text.startswith("İsmet Kaptan Mah.") or text.startswith("Konak/İZMİR"):
            run = p.add_run(text)
            p.paragraph_format.left_indent = Cm(3.75)
        else:
            run = p.add_run(text)

    # Dosyayı kaydet
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path